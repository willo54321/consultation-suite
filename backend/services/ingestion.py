import os
import uuid
from typing import Optional, List, Tuple
from pathlib import Path
import PyPDF2
from docx import Document as DocxDocument
from sqlalchemy.orm import Session

from database import Document, Chunk
from services.embeddings import generate_embedding, generate_embeddings_batch
from utils.text_processing import clean_text, split_into_chunks, count_tokens
from config import get_settings

settings = get_settings()


def process_document(
    db: Session,
    document: Document,
    file_path: str,
) -> Tuple[bool, Optional[str]]:
    """
    Process an uploaded document: extract text, chunk, and generate embeddings.
    Returns (success, error_message).
    """
    try:
        print(f"[PROCESS] Starting extraction for {file_path}", flush=True)

        # Extract text based on file type
        if document.file_type == "pdf":
            text, page_count = extract_pdf_text(file_path)
            document.page_count = page_count
        elif document.file_type == "docx":
            text = extract_docx_text(file_path)
            document.page_count = None
        elif document.file_type == "txt":
            text = extract_txt_text(file_path)
            document.page_count = None
        else:
            return False, f"Unsupported file type: {document.file_type}"

        print(f"[PROCESS] Extracted {len(text)} chars, {document.page_count} pages", flush=True)

        if not text or len(text.strip()) < 50:
            return False, "Document contains insufficient text content"

        # Clean the text
        cleaned_text = clean_text(text)
        print(f"[PROCESS] Cleaned text: {len(cleaned_text)} chars", flush=True)

        # Split into chunks
        print(f"[PROCESS] About to split into chunks...", flush=True)
        chunks = split_into_chunks(
            cleaned_text,
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
        )
        print(f"[PROCESS] Split complete", flush=True)

        if not chunks:
            return False, "Failed to create chunks from document"

        print(f"[PROCESS] Created {len(chunks)} chunks", flush=True)

        # Generate embeddings for all chunks
        chunk_texts = [chunk[0] for chunk in chunks]
        print(f"[PROCESS] Starting embedding generation for {len(chunk_texts)} chunks...", flush=True)
        embeddings = generate_embeddings_batch(chunk_texts)
        print(f"[PROCESS] Generated {len(embeddings)} embeddings", flush=True)

        # Store chunks in database
        for idx, ((chunk_text, token_count), embedding) in enumerate(zip(chunks, embeddings)):
            chunk = Chunk(
                id=uuid.uuid4(),
                document_id=document.id,
                project_id=document.project_id,
                content=chunk_text,
                chunk_index=idx,
                token_count=token_count,
                embedding=embedding,
                chunk_metadata={
                    "source_file": document.original_filename or document.filename,
                    "content_type": document.content_type,
                    "tags": document.tags or [],
                },
                source_type="document",
            )
            db.add(chunk)

        print(f"[PROCESS] Committing to database...", flush=True)

        # Mark document as processed
        document.processed = True
        document.processing_error = None
        db.commit()

        print(f"[PROCESS] Document processed successfully!", flush=True)
        return True, None

    except Exception as e:
        print(f"[PROCESS] Error: {e}", flush=True)
        document.processing_error = str(e)
        db.commit()
        return False, str(e)


def extract_pdf_text(file_path: str) -> Tuple[str, int]:
    """Extract text from PDF file."""
    text_parts = []
    page_count = 0

    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        page_count = len(reader.pages)

        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(f"[Page {page_num + 1}]\n{page_text}")

    return "\n\n".join(text_parts), page_count


def extract_docx_text(file_path: str) -> str:
    """Extract text from DOCX file."""
    doc = DocxDocument(file_path)
    paragraphs = []

    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs)


def extract_txt_text(file_path: str) -> str:
    """Extract text from TXT file."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def process_raw_text(
    db: Session,
    project_id: uuid.UUID,
    text: str,
    source_name: str = "Manual Entry",
    content_type: str = "public",
    tags: List[str] = None,
) -> Tuple[bool, Optional[str]]:
    """Process raw text content (for manual entries or scraped content)."""
    try:
        cleaned_text = clean_text(text)

        if len(cleaned_text) < 50:
            return False, "Text content too short"

        chunks = split_into_chunks(
            cleaned_text,
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
        )

        if not chunks:
            return False, "Failed to create chunks"

        chunk_texts = [chunk[0] for chunk in chunks]
        embeddings = generate_embeddings_batch(chunk_texts)

        for idx, ((chunk_text, token_count), embedding) in enumerate(zip(chunks, embeddings)):
            chunk = Chunk(
                id=uuid.uuid4(),
                project_id=project_id,
                content=chunk_text,
                chunk_index=idx,
                token_count=token_count,
                embedding=embedding,
                chunk_metadata={
                    "source_name": source_name,
                    "content_type": content_type,
                    "tags": tags or [],
                },
                source_type="manual",
            )
            db.add(chunk)

        db.commit()
        return True, None

    except Exception as e:
        return False, str(e)
